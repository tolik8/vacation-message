"""
Повідомлення про відпустку - кадри (Шамрик)
"""

import datetime
import os
import csv
# pip install python-docx
from docx import Document

# константи, початкові змінні
date_format = '%d.%m.%Y'
data_list = 'list.csv'
template = 'template.docx'

# очищаємо екран
os.system('cls' if os.name == 'nt' else 'clear')

# отримуємо робочий каталог
folder_path = os.path.dirname(os.path.realpath(__file__))
data_path = folder_path + os.sep + 'data' + os.sep
result_path = folder_path + os.sep + 'result' + os.sep

if not os.path.exists(result_path):
    os.makedirs(result_path, exist_ok=True)

# видаляємо всі файли з папки result
[os.remove(os.path.join(result_path, i)) for i in os.listdir(result_path)]

# читаємо CSV файл
with open(data_path + data_list, newline='') as csv_file:
    context = csv.reader(csv_file, delimiter=';')
    lines = [i for i in context]

pib_remember = ''
position_remember = ''

for line in lines:
    if line[1] != '':
        pib = line[1]
        pib_remember = pib
    else:
        pib = pib_remember

    if line[2] != '':
        position = line[2]
        position_remember = position
    else:
        position = position_remember

    # присвоюємо значення
    date2 = line[0]
    date1_obj = datetime.datetime.strptime(date2, date_format) - datetime.timedelta(weeks=2)
    date1 = date1_obj.strftime(date_format)
    date1_values = date1.strip().split('.')
    mydate1 = date1_values[2] + '-' + date1_values[1] + '-' + date1_values[0]

    date2_values = date2.strip().split('.')
    mydate2 = date2_values[2] + '-' + date2_values[1] + '-' + date2_values[0]

    pib_values = pib.strip().split(' ')
    pib1 = pib_values[0]
    pib2_short = pib_values[1][0]
    pib3_short = pib_values[2][0]
    pib_short = pib1 + ' ' + pib2_short + '.' + pib3_short + '.'

    # відкриваємо шаблон документа Word
    document = Document(data_path + template)

    # проходимо по всіх параграфах документу
    for paragraph in document.paragraphs:
        # замінюємо частину тексту, якщо вона містить потрібне значення
        if '%DATE2%' in paragraph.text:
            paragraph.text = paragraph.text.replace('%DATE2%', date2)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace("%NAME%", pib_short)
                cell.text = cell.text.replace("%POSITION%", position)
                cell.text = cell.text.replace("%DATE1%", date1)

    # зберігаємо новий документ Word зі значеннями з рядка
    document_name = f'result{os.sep}{mydate1} {pib1} {pib2_short}{pib3_short} {mydate2}.docx'
    print(document_name)
    document.save(document_name)
